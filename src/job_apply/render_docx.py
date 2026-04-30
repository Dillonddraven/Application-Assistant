"""DOCX rendering via python-docx.

ATS-friendly: single column, real headings, no images/text-boxes/columns. Same
visual restraint as the PDF (sans-serif body, navy accent on name + section
headers), but built from python-docx primitives so applicant tracking systems
parse the text correctly.
"""
from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

from .profile_loader import Profile, Secrets
from .render import _scrub_source_ids, collect_known_ids


NAVY = RGBColor(0x1A, 0x3A, 0x5C)
GREY = RGBColor(0x5B, 0x60, 0x68)
INK = RGBColor(0x1D, 0x1F, 0x24)


def _set_run(run, *, size_pt: float, color: RGBColor = INK,
             bold: bool = False, italic: bool = False,
             family: str = "Helvetica Neue") -> None:
    run.font.name = family
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def _set_para_spacing(para, *, before_pt: float = 0, after_pt: float = 0,
                      line_pct: int | None = None) -> None:
    fmt = para.paragraph_format
    fmt.space_before = Pt(before_pt)
    fmt.space_after = Pt(after_pt)
    if line_pct is not None:
        fmt.line_spacing = line_pct / 100.0


def _section_header(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _set_para_spacing(p, before_pt=10, after_pt=2)
    r = p.add_run(text.upper())
    _set_run(r, size_pt=9.5, color=NAVY, bold=True)
    r.font.all_caps = True
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A3A5C")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    _set_para_spacing(p, before_pt=0, after_pt=2, line_pct=120)
    p.paragraph_format.left_indent = Inches(0.18)
    r = p.runs[0] if p.runs else p.add_run("")
    r.text = text
    _set_run(r, size_pt=10.5, color=INK)


def _two_col_line(doc: Document, *, left: str, right: str,
                  left_size: float = 11.0, left_bold: bool = False,
                  left_color: RGBColor = INK, right_size: float = 9.5,
                  right_color: RGBColor = GREY) -> None:
    p = doc.add_paragraph()
    _set_para_spacing(p, before_pt=0, after_pt=1, line_pct=115)
    fmt = p.paragraph_format
    fmt.tab_stops.add_tab_stop(Inches(7.1), WD_TAB_ALIGNMENT.RIGHT)
    rl = p.add_run(left)
    _set_run(rl, size_pt=left_size, color=left_color, bold=left_bold)
    if right:
        p.add_run("\t")
        rr = p.add_run(right)
        _set_run(rr, size_pt=right_size, color=right_color)


def _fmt_period(start: str, end: str) -> str:
    end = end or "Present"
    return f"{start} – {end}" if start else end


def _contact_line_text(secrets: Secrets) -> str:
    pm = secrets.placeholder_map()
    city = pm.get("{{address_city}}") or ""
    state = pm.get("{{address_state}}") or ""
    bits = []
    if city and state:
        bits.append(f"{city}, {state}")
    elif city:
        bits.append(city)
    bits += [pm.get("{{email}}", ""), pm.get("{{phone}}", ""),
             pm.get("{{linkedin_url}}", ""), pm.get("{{github_url}}", "")]
    return "  •  ".join(b for b in bits if b)


def render_resume_docx(*, profile: Profile, tailored: dict[str, Any],
                       secrets: Secrets, out_path: Path) -> Path:
    known = collect_known_ids(profile.data)
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    style = doc.styles["Normal"]
    style.font.name = "Helvetica Neue"
    style.font.size = Pt(10.5)

    p = doc.add_paragraph()
    _set_para_spacing(p, before_pt=0, after_pt=1)
    r = p.add_run(profile.data.get("identity", {}).get("full_name", ""))
    _set_run(r, size_pt=22, color=NAVY, bold=True)

    contact = _contact_line_text(secrets)
    if contact:
        p = doc.add_paragraph()
        _set_para_spacing(p, before_pt=0, after_pt=12)
        r = p.add_run(contact)
        _set_run(r, size_pt=9.5, color=GREY)

    summary = _scrub_source_ids(tailored.get("summary") or "", known).strip()
    if summary:
        _section_header(doc, "Summary")
        p = doc.add_paragraph()
        _set_para_spacing(p, before_pt=2, after_pt=2, line_pct=130)
        r = p.add_run(summary)
        _set_run(r, size_pt=10.5, color=INK)

    skills = tailored.get("skills_emphasis") or []
    if skills:
        _section_header(doc, "Skills")
        # Two-column visual via tab stops, but text is linear (ATS-safe).
        col_count = 2
        rows = (len(skills) + col_count - 1) // col_count
        cols = [skills[i*rows:(i+1)*rows] for i in range(col_count)]
        max_rows = max(len(c) for c in cols)
        for row_i in range(max_rows):
            p = doc.add_paragraph()
            _set_para_spacing(p, before_pt=0, after_pt=1, line_pct=115)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(3.5))
            left = cols[0][row_i] if row_i < len(cols[0]) else ""
            right = cols[1][row_i] if row_i < len(cols[1]) else ""
            r = p.add_run(f"·  {left}" if left else "")
            _set_run(r, size_pt=10, color=INK)
            if right:
                p.add_run("\t")
                r2 = p.add_run(f"·  {right}")
                _set_run(r2, size_pt=10, color=INK)

    exp_id_set = {e.get("id") for e in (profile.data.get("experience") or []) if isinstance(e, dict)}
    cited_exp_ids: set[str] = set()
    bullets_by_exp: dict[str, list[str]] = {}
    for b in tailored.get("bullets") or []:
        sid = b.get("source_id", "")
        bits = sid.split(".")
        exp_id = None
        for candidate in bits[:2]:
            if candidate in exp_id_set:
                exp_id = candidate
                break
        if exp_id:
            cited_exp_ids.add(exp_id)
            bullets_by_exp.setdefault(exp_id, []).append(_scrub_source_ids(b.get("text", ""), known))

    if cited_exp_ids:
        _section_header(doc, "Experience")
        for exp in profile.data.get("experience") or []:
            if exp.get("id") not in cited_exp_ids:
                continue
            company = exp.get("company") or ""
            title = exp.get("title") or ""
            location = exp.get("location") or ""
            period = _fmt_period(exp.get("start") or "", exp.get("end") or "")
            right_meta = " • ".join(b for b in [period, location] if b)
            # Title line (company bold, role lighter, dates/location right)
            p = doc.add_paragraph()
            _set_para_spacing(p, before_pt=8, after_pt=1)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(7.1), WD_TAB_ALIGNMENT.RIGHT)
            r1 = p.add_run(company)
            _set_run(r1, size_pt=11, color=INK, bold=True)
            if title:
                r2 = p.add_run(f", {title}")
                _set_run(r2, size_pt=11, color=GREY)
            if right_meta:
                p.add_run("\t")
                r3 = p.add_run(right_meta)
                _set_run(r3, size_pt=9.5, color=GREY)
            for txt in bullets_by_exp.get(exp.get("id"), []):
                _bullet(doc, txt)

    edu = profile.data.get("education") or []
    if edu:
        _section_header(doc, "Education")
        for e in edu:
            degree = e.get("degree") or ""
            school = e.get("school") or ""
            status = e.get("status") or ""
            expected = e.get("expected_end") or ""
            right_bits: list[str] = []
            if status == "in_progress":
                right_bits.append("in progress")
                if expected:
                    right_bits.append(f"expected {expected}")
            elif e.get("end"):
                right_bits.append(str(e.get("end")))
            _two_col_line(
                doc,
                left=f"{degree}, {school}",
                right=", ".join(right_bits),
                left_bold=True,
            )

    certs = profile.data.get("certifications") or []
    if certs:
        _section_header(doc, "Certifications")
        for c in certs:
            issued = c.get("issued") or ""
            right = issued if issued and not str(issued).startswith("<") else ""
            _two_col_line(doc, left=c.get("name") or "", right=right, left_bold=False, left_size=10.5)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def render_cover_letter_docx(*, profile: Profile, tailored: dict[str, Any],
                             secrets: Secrets, out_path: Path,
                             company: str, salutation: str = "Dear Hiring Team,",
                             date_iso: str | None = None) -> Path:
    known = collect_known_ids(profile.data)
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
    style = doc.styles["Normal"]
    style.font.name = "Helvetica Neue"
    style.font.size = Pt(11)

    pm = secrets.placeholder_map()
    addr_lines: list[str] = []
    city_state_zip = " ".join(
        b for b in [
            (pm.get("{{address_city}}") or "") + ("," if pm.get("{{address_city}}") else ""),
            pm.get("{{address_state}}") or "",
            pm.get("{{address_zip}}") or "",
        ] if b.strip()
    ).strip()
    if pm.get("{{address_street}}"):
        addr_lines.append(pm["{{address_street}}"])
    if city_state_zip:
        addr_lines.append(city_state_zip)
    for line in addr_lines:
        p = doc.add_paragraph()
        _set_para_spacing(p, before_pt=0, after_pt=0)
        r = p.add_run(line)
        _set_run(r, size_pt=10, color=GREY)

    p = doc.add_paragraph()
    _set_para_spacing(p, before_pt=12, after_pt=14)
    r = p.add_run(date_iso or date.today().isoformat())
    _set_run(r, size_pt=10, color=INK)

    p = doc.add_paragraph()
    _set_para_spacing(p, before_pt=0, after_pt=10)
    r = p.add_run(salutation)
    _set_run(r, size_pt=11, color=INK, bold=False)

    for para in tailored.get("cover_letter_paragraphs") or []:
        text = _scrub_source_ids(para.get("text", ""), known).strip()
        if not text:
            continue
        p = doc.add_paragraph()
        _set_para_spacing(p, before_pt=0, after_pt=10, line_pct=140)
        r = p.add_run(text)
        _set_run(r, size_pt=11, color=INK)

    p = doc.add_paragraph()
    _set_para_spacing(p, before_pt=12, after_pt=4)
    r = p.add_run("Sincerely,")
    _set_run(r, size_pt=11, color=INK)

    name = profile.data.get("identity", {}).get("full_name", "")
    p = doc.add_paragraph()
    _set_para_spacing(p, before_pt=0, after_pt=0)
    r = p.add_run(name)
    _set_run(r, size_pt=11, color=NAVY, bold=True)

    sig_bits = [pm.get("{{email}}", ""), pm.get("{{phone}}", "")]
    sig = "  •  ".join(b for b in sig_bits if b)
    if sig:
        p = doc.add_paragraph()
        _set_para_spacing(p, before_pt=0, after_pt=0)
        r = p.add_run(sig)
        _set_run(r, size_pt=10, color=GREY)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
